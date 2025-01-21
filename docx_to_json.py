import json
from docx import Document
import os

def docx_to_json(docx_path):
    doc = Document(docx_path)
    page_content = {
        "title": doc.paragraphs[0].text,  # i am assuming the title is the first paragraph
        "intro": "",
        "sections": [],
        "tables": [],
    }
    
    current_section = None

# process paragraphs/headings/body
    for paragraph in doc.paragraphs[1:]:  # to skip the first paragraph used as title
        text = paragraph.text.strip()
        if not text:
            continue

# id headers IOT start new sections
        if paragraph.style.name.startswith("Heading"):
            if current_section:
                page_content["sections"].append(current_section)
            current_section = {"header": text, "content": "", "subsections": []}
        
# regular text goes to intro if no section or to current section
	else:
		if current_section:
                current_section["content"] += f"{text} "
        else:
                page_content["intro"] += f"{text} "

# append the last section if any there
		if current_section:
      		page_content["sections"].append(current_section)

# process ze tables
		for table in doc.tables:
       		 table_data = []
       		 for row in table.rows:
            		row_data = [cell.text.strip() for cell in row.cells]
          	        table_data.append(row_data)
        page_content["tables"].append(table_data)

    return page_content

def process_docs(folder_path):
    pages = []
    for filename in os.listdir(folder_path):
        if filename.endswith(".docx"):
            filepath = os.path.join(folder_path, filename)
            page_data = docx_to_json(filepath)
            pages.append(page_data)
    
    with open("content.json", "w", encoding="utf-8") as f:
        json.dump(pages, f, indent=4)

    print("Converted DOCX files to content.json")

# use the function with this folder path (edit if you're using this for yourself ofc so it matches your setup)
process_docs("/Users/Ivy/Desktop/docxfilesnov")
